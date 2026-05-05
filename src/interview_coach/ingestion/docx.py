import io
import re
from zipfile import BadZipFile

from docx import Document as DocxDocument

from interview_coach.ingestion.errors import ExtractionFailed

_BLANK_LINES_RE = re.compile(r"\n\s*\n\s*\n+")


def extract_docx_text(data: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(data))
    except (BadZipFile, KeyError, ValueError) as e:
        raise ExtractionFailed(f"Failed to read DOCX: {e}") from e

    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return _BLANK_LINES_RE.sub("\n\n", "\n".join(parts)).strip()
